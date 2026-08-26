#!/usr/bin/env python3
"""
Build Live AEX walkthrough HTML + Word (.docx) from crawl-meta.json + screenshots/.

Usage:
  python3 generate-docs.py
"""
from __future__ import annotations

import html
import json
import re
from collections import OrderedDict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
META_PATH = ROOT / "crawl-meta.json"
HTML_OUT = ROOT / "WALKTHROUGH.html"
DOCX_OUT = ROOT / "Live-AEX-UI-Walkthrough.docx"

# Keep in sync with generate-readme.py DESC where possible
DESC = {
    "experience-dashboard": "Tenant home / control tower. Device health, compliance posture, open alerts, and recommended actions for the selected site.",
    "home-recent-activity": "Operational activity log — who changed automations, published surveys, pushed software, or linked profiles.",
    "automation-config": "Automation library of self-help, autoheal, and scheduler profiles. Browse remediations, open details, link scope to device groups.",
    "visual-builder": "Workflow Builder for authoring automations visually: template or custom, actions, When triggers, Matches conditions, save, then link scope.",
    "insights-campaigns": "Survey campaigns for desktop satisfaction, employee satisfaction, and sentiment.",
    "data-collection": "Data-collection profiles (performance, storage, user activity, network) for DEX and compliance telemetry.",
    "protection-policies": "Device protection profiles (registry, services, application control, paths, USB).",
    "gpo-catalog": "ADMX policy template browser for Windows desired-state policies.",
    "pending-changes": "Pending Publish — staged configuration changes per machine group before they go live.",
    "innovate-assistant": "Conversational AI Helpdesk (Haya) for plain-language estate questions.",
    "innovate-investigations": "Investigation workspace on the Haya conversation surface.",
    "innovate-anomaly": "Anomaly clusters for unusual device or experience patterns.",
    "innovate-forecasts": "Forecast instances bound to Studio-published definitions.",
    "innovate-query-builder": "Ad-hoc analytics: dimensions, measures, JSON filters, Execute.",
    "innovate-agent-workbench": "Agent chat workbench with smart recommendations from telemetry.",
    "innovate-alerts": "Alert triage surface (Haya conversation in this build).",
    "dex-dashboards": "Catalog of Digital Experience (DEX) and device reports.",
    "dex-reports": "Executive Digital Experience Overview for leadership.",
    "dex-employee": "Employee Experience story: sentiment and complaint themes.",
    "dex-device": "Device Experience story: satisfaction and boot/hardware health.",
    "dex-application": "Application Experience story: business and collaboration apps.",
    "dex-network": "Network Experience story: latency and connectivity by site.",
    "compliance-score": "Digital Experience dashboard (fleet DEX score) — not an external BI login.",
    "dex-benchmarks": "Benchmarks comparing satisfaction vs endpoint performance.",
    "compliance-trends": "Trends placeholder for historical posture reporting.",
    "census": "Device inventory (Census) — OS filters, search, export, health chips.",
    "devices-device-explorer": "Coming soon — device search + full profile explorer.",
    "actions-results": "Execution History of solutions pushed to devices (success, failed, pending).",
    "automate-remote-execution": "Remote Execution — run a linked remediation now on site, group, or device.",
    "manage-swd": "Software Distribution — install/uninstall profiles and snippets.",
    "swd-push-audit": "Software distribution Deploy queue — per-device install/uninstall audit.",
    "manage-patch": "Coming soon — Windows Update / patch tracking.",
    "issues-compliance-violations": "Violations triage (Haya conversation in this build).",
    "tenants": "Organisation tenants/sites — names, org, keys, installs, license pools.",
    "org-branding": "Coming soon — branding on the IA roadmap.",
    "admin-licensing": "License pools — entitlements and seat usage.",
    "platform-integration": "Integrations hub: connectors, API keys, catalog, event subscriptions.",
    "admin-agent-mgmt": "Agent Management — versions & rings, updates, version catalog, uploads.",
    "access-security": "Users & Access — accounts, RBAC, SSO, MFA.",
    "settings-security": "Coming soon — certificates/secrets and GDPR ops.",
    "device-classification": "Device classification groups for targeting.",
    "audit-logs": "Administration Audit hub — activity, ticket, and system events.",
    "scoring": "Score Models that feed DEX and operations scoring.",
    "account-profile": "Personal profile for the signed-in operator.",
    "account-preferences": "Theme and accent color preferences.",
    "header-help-feedback": "Help & Feedback from the header.",
    "automate-scheduling": "Scheduling hub — deployed workflows on a recurring schedule.",
    "monitoring-profiles": "Monitoring / Data Collection profiles hub.",
    "admin-deployment": "Client enrolment packages hub (placeholder until package UI ships).",
    "version-catalog": "Version catalog of client packages sites can target for updates.",
    "gpo-policies": "GPO / desired-state policy definitions list.",
    "compliance-policy-execution": "Per-device policy execution status.",
    "selfhelp-library": "Automation library filtered toward self-help profiles.",
    "autoheal-library": "Automation library filtered toward autoheal profiles.",
}

SECTIONS_ORDER = [
    "Home",
    "Automate",
    "Innovate",
    "Experience",
    "Manage",
    "Administration",
    "Settings",
    "Header",
    "DEX report views",
    "Create / edit overlays",
    "Device detail",
    "Disabled menu",
    "Workflow Builder deep",
    "Software Distribution deep",
]


def load_entries(meta: dict) -> list[dict]:
    """Prefer crawl results; overlay newer recapture pages by id/file."""
    results = list(meta.get("results") or [])
    by_id = {r.get("id"): i for i, r in enumerate(results) if r.get("id")}
    by_file = {r.get("file"): i for i, r in enumerate(results) if r.get("file")}
    for p in meta.get("pages") or []:
        pid, pfile = p.get("id"), p.get("file")
        if pid in by_id:
            results[by_id[pid]] = {**results[by_id[pid]], **p}
        elif pfile in by_file:
            results[by_file[pfile]] = {**results[by_file[pfile]], **p}
        else:
            results.append(p)
    return results


def entry_title(r: dict) -> str:
    n = r.get("notes") or {}
    h1 = (n.get("h1") or "").strip()
    label = r.get("label") or r.get("id") or "Untitled"
    if h1 and h1.lower() != label.lower():
        return f"{label} — {h1}"
    return label


def entry_desc(r: dict) -> str:
    desc = DESC.get(r.get("id") or "", "")
    if not desc and r.get("description"):
        desc = str(r["description"])
    if not desc:
        desc = ((r.get("notes") or {}).get("snippet") or "")[:280] or "(see screenshot)"
    return desc


def entry_tags(r: dict) -> list[str]:
    n = r.get("notes") or {}
    tags = []
    if n.get("comingSoon"):
        tags.append("coming soon / placeholder")
    if r.get("menuEnabled") is False:
        tags.append("menu disabled")
    if r.get("kind") == "drawer":
        tags.append("drawer")
    if not r.get("ok", True):
        tags.append("capture failed")
    return tags


def slug_sec(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")


def build_html(meta: dict, entries: list[dict]) -> str:
    by_sec: OrderedDict[str, list] = OrderedDict((s, []) for s in SECTIONS_ORDER)
    for r in entries:
        sec = r.get("section") or "Other"
        by_sec.setdefault(sec, []).append(r)

    captured = meta.get("capturedAt") or datetime.now(timezone.utc).isoformat()
    base = meta.get("base") or "http://localhost:5176"
    nav = "".join(
        f'<a href="#{html.escape(slug_sec(s))}">{html.escape(s)}</a>'
        for s, items in by_sec.items()
        if items
    )

    parts = [
        "<!DOCTYPE html>",
        '<html lang="en">',
        "<head>",
        '  <meta charset="UTF-8" />',
        '  <meta name="viewport" content="width=device-width, initial-scale=1" />',
        "  <title>Live AEX product — UI walkthrough</title>",
        "  <style>",
        """
    :root { --text:#1a2333; --muted:#5b6b82; --border:#d8e0ec; --bg:#f4f6fa; --card:#fff; --accent:#1d4ed8; --tag:#eef4fc; }
    * { box-sizing: border-box; }
    body { margin:0; font-family:"Segoe UI",system-ui,-apple-system,sans-serif; color:var(--text); background:var(--bg); line-height:1.5; }
    header { background:#fff; border-bottom:1px solid var(--border); padding:28px 32px; position:sticky; top:0; z-index:10; }
    header h1 { margin:0 0 6px; font-size:24px; }
    header p { margin:0; color:var(--muted); font-size:14px; }
    nav { display:flex; flex-wrap:wrap; gap:8px; margin-top:14px; }
    nav a { font-size:12px; text-decoration:none; color:var(--accent); background:var(--tag); padding:4px 10px; border-radius:999px; border:1px solid var(--border); }
    main { max-width:980px; margin:0 auto; padding:24px 20px 64px; }
    section { margin-bottom:40px; }
    section > h2 { font-size:18px; margin:0 0 16px; padding-bottom:8px; border-bottom:2px solid var(--accent); }
    article { background:var(--card); border:1px solid var(--border); border-radius:12px; padding:20px; margin-bottom:16px; }
    article h3 { margin:0 0 8px; font-size:17px; }
    .meta { display:flex; flex-wrap:wrap; gap:8px; margin-bottom:10px; }
    .meta span { font-size:12px; background:var(--tag); border:1px solid var(--border); border-radius:6px; padding:2px 8px; color:var(--muted); font-family:ui-monospace,Menlo,monospace; }
    .desc { margin:0 0 14px; font-size:14px; }
    figure { margin:0; border:1px solid var(--border); border-radius:10px; overflow:hidden; background:#e8eef6; }
    figure img { display:block; width:100%; height:auto; }
    figcaption { font-size:12px; color:var(--muted); padding:8px 12px; background:#fafbfd; border-top:1px solid var(--border); }
    .toc { background:var(--card); border:1px solid var(--border); border-radius:12px; padding:16px 20px; margin-bottom:28px; font-size:14px; }
        """.strip(),
        "  </style>",
        "</head>",
        "<body>",
        "  <header>",
        "    <h1>Live AEX product — UI walkthrough</h1>",
        f"    <p>Nanoheal Live AEX console · <code>{html.escape(base)}</code> · {html.escape(str(captured)[:19])} · {len(entries)} screens</p>",
        f"    <nav>{nav}</nav>",
        "  </header>",
        "  <main>",
        '    <div class="toc">',
        "      <p><strong>How to read:</strong> Each card is one navigable surface or overlay. "
        "<em>Coming soon / placeholder</em> means nav is wired but UI is limited. "
        "<em>Menu disabled</em> means the sidebar item is off but the URL still works.</p>",
        "      <p>Also available as Word: <code>Live-AEX-UI-Walkthrough.docx</code> · Markdown: <code>README.md</code></p>",
        "    </div>",
    ]

    for sec, items in by_sec.items():
        if not items:
            continue
        parts.append(f'    <section id="{html.escape(slug_sec(sec))}">')
        parts.append(f"      <h2>{html.escape(sec)}</h2>")
        for r in items:
            title = entry_title(r)
            tags = entry_tags(r)
            path_s = r.get("path") or ""
            file_s = r.get("file") or ""
            img = f"screenshots/{file_s}" if file_s else ""
            meta_bits = []
            if path_s:
                meta_bits.append(f"<span>Path: {html.escape(path_s)}</span>")
            for t in tags:
                meta_bits.append(f"<span>{html.escape(t)}</span>")
            parts.append("      <article>")
            parts.append(f"        <h3>{html.escape(title)}</h3>")
            if meta_bits:
                parts.append(f'        <div class="meta">{"".join(meta_bits)}</div>')
            parts.append(f'        <p class="desc">{html.escape(entry_desc(r))}</p>')
            if file_s and (SHOTS / file_s).is_file():
                parts.append("        <figure>")
                parts.append(
                    f'          <img src="{html.escape(img)}" alt="{html.escape(title)}" loading="lazy" />'
                )
                parts.append(f"          <figcaption>{html.escape(img)}</figcaption>")
                parts.append("        </figure>")
            elif file_s:
                parts.append(f'        <p class="desc"><em>Missing image: {html.escape(img)}</em></p>')
            parts.append("      </article>")
        parts.append("    </section>")

    parts += ["  </main>", "</body>", "</html>"]
    return "\n".join(parts) + "\n"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Calibri")
    if color:
        run.font.color.rgb = color


def build_docx(meta: dict, entries: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Live AEX product — UI walkthrough")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x1A, 0x23, 0x33))

    sub = doc.add_paragraph()
    r = sub.add_run(
        f"Nanoheal Live AEX · {meta.get('base', 'http://localhost:5176')} · "
        f"{str(meta.get('capturedAt', ''))[:19]} · {len(entries)} screens"
    )
    set_run_font(r, size=10, color=RGBColor(0x5B, 0x6B, 0x82))

    intro = doc.add_paragraph()
    r = intro.add_run(
        "Each entry is one navigable console surface or overlay with a screenshot. "
        "Coming soon / menu-disabled items are still included when the URL is reachable."
    )
    set_run_font(r, size=11)

    by_sec: OrderedDict[str, list] = OrderedDict((s, []) for s in SECTIONS_ORDER)
    for r in entries:
        sec = r.get("section") or "Other"
        by_sec.setdefault(sec, []).append(r)

    for sec, items in by_sec.items():
        if not items:
            continue
        h = doc.add_heading(sec, level=1)
        for run in h.runs:
            set_run_font(run, size=16, bold=True, color=RGBColor(0x1D, 0x4E, 0xD8))

        for ent in items:
            h2 = doc.add_heading(entry_title(ent), level=2)
            for run in h2.runs:
                set_run_font(run, size=13, bold=True)

            meta_line = []
            if ent.get("path"):
                meta_line.append(f"Path: {ent['path']}")
            tags = entry_tags(ent)
            if tags:
                meta_line.append(" · ".join(tags))
            if meta_line:
                p = doc.add_paragraph()
                r = p.add_run(" · ".join(meta_line))
                set_run_font(r, size=9, color=RGBColor(0x5B, 0x6B, 0x82))

            p = doc.add_paragraph()
            r = p.add_run(entry_desc(ent))
            set_run_font(r, size=11)

            file_s = ent.get("file") or ""
            img_path = SHOTS / file_s if file_s else None
            if img_path and img_path.is_file():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.8))
                    last = doc.paragraphs[-1]
                    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph()
                    cr = cap.add_run(f"screenshots/{file_s}")
                    set_run_font(cr, size=8, color=RGBColor(0x5B, 0x6B, 0x82))
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    p = doc.add_paragraph()
                    r = p.add_run(f"(Could not embed image {file_s}: {e})")
                    set_run_font(r, size=9, color=RGBColor(0x9A, 0x34, 0x12))

    doc.save(DOCX_OUT)


def main() -> None:
    meta = json.loads(META_PATH.read_text())
    entries = load_entries(meta)
    HTML_OUT.write_text(build_html(meta, entries), encoding="utf-8")
    build_docx(meta, entries)
    print(f"Wrote {HTML_OUT.name} ({HTML_OUT.stat().st_size // 1024} KB)")
    print(f"Wrote {DOCX_OUT.name} ({DOCX_OUT.stat().st_size // 1024} KB)")
    print(f"Entries: {len(entries)}")


if __name__ == "__main__":
    main()
